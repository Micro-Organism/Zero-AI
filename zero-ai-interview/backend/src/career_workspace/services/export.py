import json
from io import BytesIO

from docx import Document


def resume_markdown(snapshot: dict) -> str:
    resume = snapshot.get("resume", snapshot)
    lines = [f"# {resume.get('title', '简历')}", ""]
    if resume.get("summary"):
        lines.extend([resume["summary"], ""])
    for work in snapshot.get("work_experiences", []):
        lines.extend([f"## {work['company']} · {work['role']}", work.get("description", ""), ""])
        lines.extend(f"- {item}" for item in work.get("achievements", []))
        lines.append("")
    projects = snapshot.get("project_experiences", [])
    if projects:
        lines.append("## 项目经历")
    for project in projects:
        lines.extend([f"### {project['name']}", project.get("background", "")])
        lines.extend(f"- {item}" for item in project.get("achievements", []))
        lines.append("")
    skills = snapshot.get("skills", [])
    if skills:
        lines.append("## 技能")
        lines.extend(f"- {item['name']}：{item.get('evidence', '')}" for item in skills)
    return "\n".join(lines).strip() + "\n"


def resume_docx(snapshot: dict) -> bytes:
    document = Document()
    resume = snapshot.get("resume", snapshot)
    document.add_heading(resume.get("title", "简历"), level=0)
    if resume.get("summary"):
        document.add_paragraph(resume["summary"])
    for work in snapshot.get("work_experiences", []):
        document.add_heading(f"{work['company']} · {work['role']}", level=1)
        document.add_paragraph(work.get("description", ""))
        for item in work.get("achievements", []):
            document.add_paragraph(item, style="List Bullet")
    for project in snapshot.get("project_experiences", []):
        document.add_heading(project["name"], level=1)
        document.add_paragraph(project.get("background", ""))
        for item in project.get("achievements", []):
            document.add_paragraph(item, style="List Bullet")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def resume_json(snapshot: dict) -> bytes:
    return json.dumps(snapshot, ensure_ascii=False, indent=2).encode()
